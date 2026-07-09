from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Add Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("PHỤ LỤC XII\n")
run.bold = True
run = p.add_run("MẪU BÁO CÁO TỔNG HỢP TÌNH HÌNH TAI NẠN LAO ĐỘNG CẤP CƠ SỞ (6 THÁNG HOẶC CẢ NĂM)\n")
run.bold = True
run = p.add_run("(Kèm theo Nghị định số 39/2016/NĐ-CP ngày 15 tháng 5 năm 2016 của Chính phủ)\n\n")

# Header info
p = doc.add_paragraph()
p.add_run("Đơn vị báo cáo: ").bold = True
p.add_run("{companyName}\n")
p.add_run("BÁO CÁO TỔNG HỢP TÌNH HÌNH TAI NẠN LAO ĐỘNG\n").bold = True
p.add_run("Kỳ báo cáo ({periodName}) năm {reportYear}\n")
p.add_run("Ngày báo cáo: {reportDate}\n")
p.add_run("Thuộc loại hình cơ sở (doanh nghiệp): {companyType}    Mã loại hình cơ sở: {typeCode}\n")
p.add_run("Đơn vị nhận báo cáo: Sở Lao động - Thương binh và Xã hội.\n")
p.add_run("Lĩnh vực sản xuất chính của cơ sở: {companyField} Mã lĩnh vực: {fieldCode}\n")
p.add_run("Tổng số lao động của cơ sở: {totalEmployees} người, trong đó nữ: {femaleEmployees} người\n")
p.add_run("Tổng quỹ lương: {totalSalary} triệu đồng\n\n")

doc.add_paragraph("I. Tình hình chung tai nạn lao động").bold = True

# TABLE 1
table1 = doc.add_table(rows=4, cols=13)
table1.style = 'Table Grid'
table1.alignment = WD_TABLE_ALIGNMENT.CENTER

def set_cell_text(cell, text, bold=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = bold

# Merge header cells
table1.cell(0,0).merge(table1.cell(2,0))
set_cell_text(table1.cell(0,0), "Loại hình cơ sở", bold=True)

table1.cell(0,1).merge(table1.cell(2,1))
set_cell_text(table1.cell(0,1), "Mã số", bold=True)

table1.cell(0,2).merge(table1.cell(0,3))
set_cell_text(table1.cell(0,2), "Cơ sở", bold=True)

table1.cell(0,4).merge(table1.cell(0,6))
set_cell_text(table1.cell(0,4), "Lực lượng lao động", bold=True)

table1.cell(0,7).merge(table1.cell(0,9))
set_cell_text(table1.cell(0,7), "Tổng số tai nạn lao động", bold=True)

table1.cell(0,10).merge(table1.cell(0,11))
set_cell_text(table1.cell(0,10), "Tần suất TNLĐ", bold=True)

table1.cell(0,12).merge(table1.cell(2,12))
set_cell_text(table1.cell(0,12), "Ghi chú", bold=True)

# Row 1 merges
table1.cell(1,2).merge(table1.cell(2,2))
set_cell_text(table1.cell(1,2), "Tổng số", bold=True)

table1.cell(1,3).merge(table1.cell(2,3))
set_cell_text(table1.cell(1,3), "Số cơ sở tham gia", bold=True)

table1.cell(1,4).merge(table1.cell(2,4))
set_cell_text(table1.cell(1,4), "Tổng số lao động", bold=True)

table1.cell(1,5).merge(table1.cell(2,5))
set_cell_text(table1.cell(1,5), "Số LĐ của cơ sở tham gia BC", bold=True)

table1.cell(1,6).merge(table1.cell(2,6))
set_cell_text(table1.cell(1,6), "Số LĐ nữ", bold=True)

table1.cell(1,7).merge(table1.cell(1,9))
set_cell_text(table1.cell(1,7), "Số người bị TNLĐ", bold=True)

table1.cell(1,10).merge(table1.cell(2,10))
set_cell_text(table1.cell(1,10), "KTNLĐ", bold=True)

table1.cell(1,11).merge(table1.cell(2,11))
set_cell_text(table1.cell(1,11), "KChết", bold=True)

# Row 2 merges
set_cell_text(table1.cell(2,7), "Tổng số", bold=True)
set_cell_text(table1.cell(2,8), "Số người chết", bold=True)
set_cell_text(table1.cell(2,9), "Số người bị thương nặng", bold=True)

# Data tags for Table 1
row = table1.rows[3]
set_cell_text(row.cells[0], "Tổng số", bold=True)
set_cell_text(row.cells[1], "")
set_cell_text(row.cells[2], "{t1_c1}")
set_cell_text(row.cells[3], "{t1_c2}")
set_cell_text(row.cells[4], "{t1_c3}")
set_cell_text(row.cells[5], "{t1_c4}")
set_cell_text(row.cells[6], "{t1_c5}")
set_cell_text(row.cells[7], "{t1_c6}")
set_cell_text(row.cells[8], "{t1_c7}")
set_cell_text(row.cells[9], "{t1_c8}")
set_cell_text(row.cells[10], "{t1_c9}")
set_cell_text(row.cells[11], "{t1_c10}")
set_cell_text(row.cells[12], "")

def add_lh_row(table, name, idx):
    row = table.add_row()
    set_cell_text(row.cells[0], name)
    set_cell_text(row.cells[1], str(idx))
    for c in range(2, 12):
        set_cell_text(row.cells[c], "{lh" + str(idx) + "_c" + str(c-1) + "}")
    set_cell_text(row.cells[12], "")

add_lh_row(table1, "Doanh nghiệp nhà nước", 2)
add_lh_row(table1, "Doanh nghiệp có vốn góp của NN", 3)
add_lh_row(table1, "Công ty trách nhiệm hữu hạn", 4)
add_lh_row(table1, "Công ty cổ phần", 5)
add_lh_row(table1, "Doanh nghiệp có vốn ĐT nước ngoài", 6)
add_lh_row(table1, "Doanh nghiệp tư nhân", 7)
add_lh_row(table1, "Cơ quan HCSN, Đảng, đoàn thể", 8)
add_lh_row(table1, "Hợp tác xã", 9)
add_lh_row(table1, "Khác", 10)


doc.add_paragraph("\nII. Phân loại TNLĐ").bold = True

table2 = doc.add_table(rows=4, cols=15)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header structure for Table 2
table2.cell(0,0).merge(table2.cell(2,0))
set_cell_text(table2.cell(0,0), "Tên chỉ tiêu thống kê", bold=True)

table2.cell(0,1).merge(table2.cell(2,1))
set_cell_text(table2.cell(0,1), "Mã số", bold=True)

table2.cell(0,2).merge(table2.cell(0,8))
set_cell_text(table2.cell(0,2), "Phân loại TNLĐ theo mức độ thương tật", bold=True)

table2.cell(0,9).merge(table2.cell(0,14))
set_cell_text(table2.cell(0,9), "Thiệt hại do TNLĐ", bold=True)

# Row 1 merges
table2.cell(1,2).merge(table2.cell(1,4))
set_cell_text(table2.cell(1,2), "Số vụ TNLĐ", bold=True)

table2.cell(1,5).merge(table2.cell(1,8))
set_cell_text(table2.cell(1,5), "Số người bị nạn (Người)", bold=True)

table2.cell(1,9).merge(table2.cell(2,9))
set_cell_text(table2.cell(1,9), "Tổng số ngày nghỉ vì TNLĐ", bold=True)

table2.cell(1,10).merge(table2.cell(2,10))
set_cell_text(table2.cell(1,10), "Tổng số tiền (1.000 đ)", bold=True)

table2.cell(1,11).merge(table2.cell(1,13))
set_cell_text(table2.cell(1,11), "Chi tiết", bold=True)

table2.cell(1,14).merge(table2.cell(2,14))
set_cell_text(table2.cell(1,14), "Thiệt hại tài sản", bold=True)

# Row 2 merges
set_cell_text(table2.cell(2,2), "Tổng số", bold=True)
set_cell_text(table2.cell(2,3), "Số vụ có người chết", bold=True)
set_cell_text(table2.cell(2,4), "Số vụ có từ 2 người bị nạn trở lên", bold=True)

set_cell_text(table2.cell(2,5), "Tổng số", bold=True)
set_cell_text(table2.cell(2,6), "Số LĐ nữ", bold=True)
set_cell_text(table2.cell(2,7), "Số người chết", bold=True)
set_cell_text(table2.cell(2,8), "Số người bị thương nặng", bold=True)

set_cell_text(table2.cell(2,11), "Y Tế", bold=True)
set_cell_text(table2.cell(2,12), "Trả lương", bold=True)
set_cell_text(table2.cell(2,13), "Bồi thường", bold=True)

# Row 3 (Tổng số TNLĐ)
row = table2.rows[3]
set_cell_text(row.cells[0], "Tổng số", bold=True)
set_cell_text(row.cells[1], "")
for c in range(2, 15):
    set_cell_text(row.cells[c], "{t2_c" + str(c+1) + "}")

def add_t2_row(table, title, code, prefix):
    row = table.add_row()
    set_cell_text(row.cells[0], title)
    set_cell_text(row.cells[1], code)
    for c in range(2, 15):
        set_cell_text(row.cells[c], "{" + prefix + "_c" + str(c+1) + "}")
    return row

start_r = len(table2.rows)
add_t2_row(table2, "Phân theo nguyên nhân", "1", "r8")
add_t2_row(table2, "Phân theo nguyên nhân", "2", "r9")
add_t2_row(table2, "Phân theo nguyên nhân", "3", "r10")
add_t2_row(table2, "Phân theo nguyên nhân", "4", "r11")
add_t2_row(table2, "Phân theo nguyên nhân", "5", "r12")
add_t2_row(table2, "Phân theo nguyên nhân", "6", "r13")
add_t2_row(table2, "Phân theo nguyên nhân", "7", "r15")
add_t2_row(table2, "Phân theo nguyên nhân", "8", "r16")
add_t2_row(table2, "Phân theo nguyên nhân", "9", "r17")
table2.cell(start_r, 0).merge(table2.cell(start_r + 8, 0))

# Factors
row = table2.add_row()
set_cell_text(row.cells[0], "{#factors}Phân theo yếu tố chấn thương")
set_cell_text(row.cells[1], "{code}")
for c in range(2, 15):
    set_cell_text(row.cells[c], "{c" + str(c+1) + "}")
set_cell_text(row.cells[14], "{c15}\n{/factors}")

# Occupations
row = table2.add_row()
set_cell_text(row.cells[0], "{#occupations}Phân theo ngành nghề")
set_cell_text(row.cells[1], "{code}")
for c in range(2, 15):
    set_cell_text(row.cells[c], "{c" + str(c+1) + "}")
set_cell_text(row.cells[14], "{c15}\n{/occupations}")

add_t2_row(table2, "TNLĐ được hưởng chế độ", "201", "t2_201")
add_t2_row(table2, "Tổng số (3=1+2)", "301", "t2_301")


doc.add_paragraph("\nĐẠI DIỆN NGƯỜI SỬ DỤNG LAO ĐỘNG").alignment = WD_ALIGN_PARAGRAPH.RIGHT
p = doc.add_paragraph("(Ký, ghi rõ họ tên, chức vụ, đóng dấu)")
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.save("d:/VNA_Group_Intern_Nhom2/fe/public/template_summary.docx")
print("Saved to d:/VNA_Group_Intern_Nhom2/fe/public/template_summary.docx")
