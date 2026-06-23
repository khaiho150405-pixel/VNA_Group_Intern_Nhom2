import docx

template_path = 'd:\\VNA_Group_Intern_Nhom2\\template.docx'
output_path = 'd:\\VNA_Group_Intern_Nhom2\\fe\\public\\template.docx'

doc = docx.Document(template_path)

doc.paragraphs[2].runs[0].text = 'Đơn vị báo cáo: {companyName}'

doc.paragraphs[3].runs[4].text = 'Kỳ báo cáo {periodName} năm {reportYear}'
doc.paragraphs[3].runs[7].text = '\xa0{reportDate}'

doc.paragraphs[5].runs[1].text = '{totalEmployees} '
doc.paragraphs[5].runs[3].text = '{femaleEmployees} '

doc.paragraphs[6].runs[1].text = '{totalSalary} '

t0 = doc.tables[0]
t0.cell(0, 0).paragraphs[0].runs[0].text = 'Địa chỉ: {companyAddress}'
t0.cell(0, 1).paragraphs[0].runs[2].text = ': {addressCode}'

t1 = doc.tables[1]
t1.cell(0, 0).paragraphs[0].runs[3].text = '{companyType}    '
t1.cell(0, 0).paragraphs[0].runs[4].text = 'Mã loại hình cơ sở: {typeCode}'

t2 = doc.tables[2]
t2.cell(0, 0).paragraphs[0].runs[1].text = '{companyField}'
t2.cell(0, 0).paragraphs[0].runs[3].text = ''
t2.cell(0, 0).paragraphs[0].runs[4].text = '\xa0Mã lĩnh vực: {fieldCode}'
t5 = doc.tables[5]
p5 = t5.cell(0, 1).paragraphs[0]
p5.runs[2].text = '(K'
p5.runs[3].text = 'ý, ghi rõ họ tên, chức vụ'
p5.runs[4].text = ', đóng dấu)'

t3 = doc.tables[3]
def insert_tags(row_idx, prefix):
    row = t3.rows[row_idx]
    for i in range(2, 13):
        if row.cells[i].paragraphs[0].runs:
            row.cells[i].paragraphs[0].runs[0].text = f'{{{prefix}_c{i+1}}}'
            for r in row.cells[i].paragraphs[0].runs[1:]: r.text = ''
        else:
            row.cells[i].text = f'{{{prefix}_c{i+1}}}'

insert_tags(5, 't1')
for i in range(8, 14): insert_tags(i, f'r{i}')
for i in range(15, 17): insert_tags(i, f'r{i}')
insert_tags(17, 'r17')

r19 = t3.rows[19]
if r19.cells[0].paragraphs[0].runs: r19.cells[0].paragraphs[0].runs[0].text = '{#factors}{name}'
else: r19.cells[0].text = '{#factors}{name}'
if r19.cells[1].paragraphs[0].runs: r19.cells[1].paragraphs[0].runs[0].text = '{code}'
else: r19.cells[1].text = '{code}'

for i in range(2, 12):
    if r19.cells[i].paragraphs[0].runs: r19.cells[i].paragraphs[0].runs[0].text = f'{{c{i+1}}}'
    else: r19.cells[i].text = f'{{c{i+1}}}'

if r19.cells[12].paragraphs[0].runs: r19.cells[12].paragraphs[0].runs[0].text = '{c13}{/factors}'
else: r19.cells[12].text = '{c13}{/factors}'

r21 = t3.rows[21]
if r21.cells[0].paragraphs[0].runs: r21.cells[0].paragraphs[0].runs[0].text = '{#occupations}{name}'
else: r21.cells[0].text = '{#occupations}{name}'
if r21.cells[1].paragraphs[0].runs: r21.cells[1].paragraphs[0].runs[0].text = '{code}'
else: r21.cells[1].text = '{code}'

for i in range(2, 12):
    if r21.cells[i].paragraphs[0].runs: r21.cells[i].paragraphs[0].runs[0].text = f'{{c{i+1}}}'
    else: r21.cells[i].text = f'{{c{i+1}}}'

if r21.cells[12].paragraphs[0].runs: r21.cells[12].paragraphs[0].runs[0].text = '{c13}{/occupations}'
else: r21.cells[12].text = '{c13}{/occupations}'

insert_tags(22, 't2')
insert_tags(23, 't3')

t4 = doc.tables[4]
if len(t4.rows) == 4:
    row4 = t4.add_row()
else:
    row4 = t4.rows[4]
for i in range(6):
    if not row4.cells[i].paragraphs[0].runs:
        row4.cells[i].text = f'{{t4_c{i+1}}}'
    else:
        row4.cells[i].paragraphs[0].runs[0].text = f'{{t4_c{i+1}}}'

doc.save(output_path)
